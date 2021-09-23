import docx

ADD_WORD_NAME = r"/diffExample.docx"


class WordWorker:
    def __init__(self, way):
        self.__way = way + ADD_WORD_NAME
        self.__doc = docx.Document()

    def insert_info(self, text, image_way='test.jpg'):
        self.__doc.add_paragraph(text)

        self.__doc.add_picture(image_way)

    def save(self):
        self.__doc.save(self.__way)

    def set_way(self, way):
        self.__way = way + ADD_WORD_NAME
